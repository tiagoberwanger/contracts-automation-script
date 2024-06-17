from docx import Document
from datetime import date

from utils import data_por_extenso, valor_formatado_por_extenso

ANO_ATUAL = date.today().year

novo_contrato = input('É um novo contrato? (S ou N):')
eh_novo_contrato = True if novo_contrato == 'S' else False

tipo_imovel = None
cpf = None
nacionalidade = None
naturalidade = None
estado_civil = None
numero = None
data_entrada = None
data_entrada_por_extenso = None
data_saida = None
data_saida_por_extenso = None
valor_formatado = None
valor_por_extenso = None
data_da_assinatura = None

if eh_novo_contrato:
    nome = input('Digite o nome do inquilino:')
    cpf = input('Digite o CPF do inquilino (formato ___.___.___.___.___):')
    nacionalidade = input('Digite sua nacionalidade:')
    naturalidade = input('Digite em qual cidade e estado você é natural (formato _______/__):')
    estado_civil = input('Digite seu estado civil:')
    tipo_imovel = input('Digite o tipo do imóvel (quarto, quitinete ou apartamento):')
    numero = input('Digite o número do imóvel:')
    data_entrada = input('Digite a data de entrada/início da vigência do contrato (formato __/__/____):')
    data_entrada_por_extenso = data_por_extenso(data_entrada)
    data_saida = input('Digite a data de saída/fim da vigência do contrato (formato __/__/____):')
    data_saida_por_extenso = data_por_extenso(data_saida)
    valor = input('Digite o valor do aluguel (apenas números):')
    valor_formatado = f'R$ {valor},00'
    valor_por_extenso = valor_formatado_por_extenso(valor)
    data_da_assinatura = input('Digite a data da assinatura (formato __/__/____):')
else:
    nome = input('Digite o nome do inquilino que você quer renovar o contrato:')
    # TODO Separar novos contratos de renovações
    print('Recurso não disponível!')

# TODO Validação dos inputs e verificação de dados presentes

dados_alterar_renovacao = {
    '{numero}': numero,
    '{data_entrada}': data_entrada,
    '{data_entrada_por_extenso}': data_entrada_por_extenso,
    '{data_saida}': data_saida,
    '{data_saida_por_extenso}': data_saida_por_extenso,
    '{valor}': valor_formatado,
    '{valor_por_extenso}': valor_por_extenso,
    '{data_da_assinatura}': data_da_assinatura
}

dados_alterar_novo = {
    **dados_alterar_renovacao,
    '{nome}': nome.upper(),
    '{cpf}': cpf,
    '{nacionalidade}': nacionalidade,
    '{naturalidade}': naturalidade,
    '{estado_civil}': estado_civil,
}


def substituir_dados(document, is_new=True):
    # DONE Pesquisar como fazer um script em py
    # DONE Pesquisar libs para manipular .docx
    # DONE Criar um método que recebe um input de texto, busca por palavras chave e substitui elas pelo texto inserido, devolve um output do texto modificado.
    data = dados_alterar_novo if is_new else dados_alterar_renovacao
    for p in document.paragraphs:
        for k, v in data.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    print("Novo contrato criado com sucesso!")


nome_formatado = nome.replace(' ', '_').lower()

# Define qual o contrato
documento_novo_contrato = Document(f'./contratos/contrato_{tipo_imovel}_modelo.docx')
documento_renovacao = Document(f'./contratos/contrato_{tipo_imovel}_modelo.docx')
# DONE Separar modelos de: quitinete, apartamento, quarto
contrato_a_atualizar = documento_novo_contrato if eh_novo_contrato else documento_renovacao

# Substitui o texto pelas informações fornecidas
substituir_dados(contrato_a_atualizar, eh_novo_contrato)

# Salva o documento modificado
contrato_a_atualizar.save(f'./contratos/contrato_{tipo_imovel}_{nome_formatado}_{ANO_ATUAL}.docx')

# TODO Criar um formulário para o usuário inserir esses dados (front)
# TODO Salvar, após concluído, esse contrato em formato PDF
# TODO Enviar, após preenchido, esse contrato para o e-mail do inquilino
